"""
run_reporter.py
---------------
Runs the 4-layer comparator then generates a full Word report.

Report sections:
  0. Cover page
  1. Executive Summary
  2. Layer 1 - Structural Validation
  3. Layer 2 - Data-Level Validation
  4. Layer 3 - Business Rule Validation
  5. Layer 4 - Statistical Validation
  6. Row-Level Differences
  7. LLM Review
  8. Appendix - Full Diff Table

Usage:
    python run_reporter.py
"""

import sys
import os
import importlib.util
import pathlib
import pandas as pd
from datetime import datetime
from loguru import logger
from dotenv import load_dotenv

load_dotenv()
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

logger.remove()
logger.add(sys.stdout,
    format="<green>{time:HH:mm:ss}</green> | <level>{level:<8}</level> | {message}",
    colorize=True, level="DEBUG")

OUTPUT_DIR = "outputs"
SEPARATOR  = "=" * 60
STEP_LINE  = "-" * 60

_root = pathlib.Path(__file__).parent


def _load_module(name, filename):
    spec = importlib.util.spec_from_file_location(name, str(_root / filename))
    mod  = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(mod)
    return mod


def run_comparison():
    comp = _load_module("run_comparator", "run_comparator.py")
    return comp.run_comparison()


# ==================================================================
# DOCX HELPERS
# ==================================================================

def _make_doc():
    from docx import Document
    from docx.shared import Inches
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)
    return doc


def _shade(row, hex_color):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    for cell in row.cells:
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  hex_color)
        tcPr.append(shd)


def _heading(doc, text, level, color):
    from docx.shared import RGBColor
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(*color)
    return h


def _status_color(passed):
    return "D5F5E3" if passed else "FADBD8"   # green / red bg


def _make_table(doc, headers, rows_data, col_statuses=None, header_bg="1A5C8A"):
    """
    Build a styled table.
    col_statuses: optional list of bool per data row to shade entire row.
    """
    from docx.shared import RGBColor
    WHITE = RGBColor(0xFF, 0xFF, 0xFF)
    BLUE  = RGBColor(0x1A, 0x5C, 0x8A)

    ncols = len(headers)
    tbl   = doc.add_table(rows=1 + len(rows_data), cols=ncols)
    tbl.style = "Table Grid"

    # Header
    hrow = tbl.rows[0]
    for j, h in enumerate(headers):
        c = hrow.cells[j]
        c.text = h
        c.paragraphs[0].runs[0].bold           = True
        c.paragraphs[0].runs[0].font.color.rgb = WHITE
    _shade(hrow, header_bg)

    # Data rows
    for i, row_vals in enumerate(rows_data):
        row = tbl.rows[i + 1]
        for j, val in enumerate(row_vals):
            row.cells[j].text = str(val) if val is not None else ""
        if col_statuses is not None:
            bg = _status_color(col_statuses[i])
            _shade(row, bg)

    return tbl


# ==================================================================
# COVER
# ==================================================================

def _section_cover(doc, results):
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    BLUE  = RGBColor(0x1A, 0x5C, 0x8A)
    GREEN = RGBColor(0x1E, 0x8B, 0x4C)
    RED   = RGBColor(0xC0, 0x39, 0x2B)
    GREY  = RGBColor(0x66, 0x66, 0x66)

    t = doc.add_heading("Migration Validation Report", 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in t.runs: r.font.color.rgb = BLUE

    for line in ["Job: JOB_NXS_T26_Payment_Status_SAP_Coupa",
                 f"Generated: {results['run_date']}"]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line)
        r.font.color.rgb = GREY
        r.font.size = Pt(11)

    doc.add_paragraph()

    decision = results["decision"]
    color    = GREEN if decision == "PASS" else RED
    badge    = doc.add_paragraph()
    badge.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = badge.add_run(f"  DECISION: {decision}  ")
    r.bold = True
    r.font.size      = Pt(20)
    r.font.color.rgb = color

    doc.add_paragraph()
    mr = results["match_rate"]
    p  = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"Match Rate: {mr:.2%}   |   Threshold: {results['threshold']:.0%}   |   "
              f"Talend rows: {results['talend_rows']}   |   SQL rows: {results['sql_rows']}")


# ==================================================================
# SECTION 1 -- EXECUTIVE SUMMARY
# ==================================================================

def _section_executive_summary(doc, results):
    BLUE  = (0x1A, 0x5C, 0x8A)
    layers = results.get("layers", {})
    l1 = layers.get("layer1") or {}
    l2 = layers.get("layer2") or {}
    l3 = layers.get("layer3") or {}
    l4 = layers.get("layer4") or {}

    _heading(doc, "1. Executive Summary", 1, BLUE)

    summary_rows = [
        ("Report Date",       results["run_date"],       None),
        ("Talend Reference",  results["talend_path"],    None),
        ("SQL Table",         results["sql_table"],      None),
        ("Final Decision",    results["decision"],       results["decision"] == "PASS"),
        ("Match Rate",        f"{results['match_rate']:.2%}", results["match_rate"] >= results["threshold"]),
        ("Pass Threshold",    f"{results['threshold']:.0%}", None),
        ("Talend Row Count",  str(results["talend_rows"]), None),
        ("SQL Row Count",     str(results["sql_rows"]),   results["talend_rows"] == results["sql_rows"]),
    ]
    _make_table(doc,
        headers=["Field", "Value"],
        rows_data=[(r[0], r[1]) for r in summary_rows],
        col_statuses=[r[2] if r[2] is not None else True for r in summary_rows]
    )
    doc.add_paragraph()

    # Layer status summary
    _heading(doc, "Layer Status Overview", 2, BLUE)
    layer_rows = [
        ("Layer 1", "Structural Validation",    l1.get("passed", False), "Fail-fast" if l1.get("fatal") else ""),
        ("Layer 2", "Data-Level Validation",    l2.get("passed", True)  if l2 else None, ""),
        ("Layer 3", "Business Rule Validation", l3.get("passed", True)  if l3 else None, ""),
        ("Layer 4", "Statistical Validation",   l4.get("passed", True)  if l4 else None, ""),
    ]
    _make_table(doc,
        headers=["Layer", "Name", "Result", "Note"],
        rows_data=[
            (r[0], r[1],
             "PASS" if r[2] else ("FAIL" if r[2] is False else "SKIPPED"),
             r[3])
            for r in layer_rows
        ],
        col_statuses=[r[2] if r[2] is not None else True for r in layer_rows]
    )
    doc.add_paragraph()

    # Fail reasons
    reasons = results.get("fail_reasons", [])
    if reasons:
        _heading(doc, "Failure Reasons", 2, BLUE)
        for reason in reasons:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(reason)
        doc.add_paragraph()


# ==================================================================
# SECTION 2 -- LAYER 1: STRUCTURAL
# ==================================================================

def _section_layer1(doc, results):
    BLUE = (0x1A, 0x5C, 0x8A)
    l1   = (results.get("layers") or {}).get("layer1") or {}
    chks = l1.get("checks", {})

    _heading(doc, "2. Layer 1 - Structural Validation", 1, BLUE)
    doc.add_paragraph(
        "Validates schema compatibility. Fail-fast: if any check fails "
        "the pipeline stops and Layers 2-4 are skipped."
    )
    doc.add_paragraph()

    # 1.1 Column naming
    _heading(doc, "2.1 Column Naming", 2, BLUE)
    c11 = chks.get("1.1_column_naming", {})
    _make_table(doc,
        headers=["Check", "Talend", "SQL", "Status"],
        rows_data=[
            ("Column count",
             str(len(c11.get("common_columns", [])) + len(c11.get("missing_in_sql", []))),
             str(len(c11.get("common_columns", [])) + len(c11.get("extra_in_sql", []))),
             "PASS" if c11.get("passed") else "FAIL"),
            ("Missing in SQL",  ", ".join(c11.get("missing_in_sql", [])) or "None", "-", "-"),
            ("Extra in SQL",    "-", ", ".join(c11.get("extra_in_sql", [])) or "None", "-"),
            ("Common columns",  str(len(c11.get("common_columns", []))), str(len(c11.get("common_columns", []))), "-"),
        ],
        col_statuses=[c11.get("passed", False), True, True, True]
    )
    doc.add_paragraph()

    # 1.2 Column count
    _heading(doc, "2.2 Column Count", 2, BLUE)
    c12 = chks.get("1.2_column_count", {})
    _make_table(doc,
        headers=["Talend Columns", "SQL Columns", "Match", "Status"],
        rows_data=[(
            str(c12.get("talend_count", "-")),
            str(c12.get("sql_count", "-")),
            "YES" if c12.get("passed") else "NO",
            "PASS" if c12.get("passed") else "FAIL"
        )],
        col_statuses=[c12.get("passed", False)]
    )
    doc.add_paragraph()

    # 1.3 Data types
    _heading(doc, "2.3 Data Types", 2, BLUE)
    c13       = chks.get("1.3_data_types", {})
    type_iss  = c13.get("type_issues", [])
    if not type_iss:
        doc.add_paragraph("All column data types are compatible.")
    else:
        _make_table(doc,
            headers=["Column", "Talend Type", "SQL Type", "Status"],
            rows_data=[(i["column"], i["talend_type"], i["sql_type"], "FAIL") for i in type_iss],
            col_statuses=[False] * len(type_iss)
        )
    doc.add_paragraph()

    # 1.4 Nullability
    _heading(doc, "2.4 Nullability", 2, BLUE)
    c14      = chks.get("1.4_nullability", {})
    null_iss = c14.get("null_issues", [])
    if not null_iss:
        doc.add_paragraph("Null ratios are consistent across all columns.")
    else:
        _make_table(doc,
            headers=["Column", "Talend Null %", "SQL Null %", "Diff %", "Status"],
            rows_data=[(i["column"], f"{i['talend_null_pct']}%",
                        f"{i['sql_null_pct']}%", f"{i['diff_pct']}%", "FAIL")
                       for i in null_iss],
            col_statuses=[False] * len(null_iss)
        )
    doc.add_paragraph()

    # 1.5 Column order
    _heading(doc, "2.5 Column Order", 2, BLUE)
    c15 = chks.get("1.5_column_order", {})
    t_order = c15.get("talend_order", [])
    s_order = c15.get("sql_order", [])
    if c15.get("passed", True):
        doc.add_paragraph("Columns are in the same order in both datasets.")
    else:
        doc.add_paragraph("Columns are in different order (warning only -- comparison proceeds).")
        max_len = max(len(t_order), len(s_order))
        _make_table(doc,
            headers=["Position", "Talend Column", "SQL Column", "Match"],
            rows_data=[
                (str(i + 1),
                 t_order[i] if i < len(t_order) else "-",
                 s_order[i] if i < len(s_order) else "-",
                 "YES" if i < len(t_order) and i < len(s_order) and t_order[i] == s_order[i] else "NO")
                for i in range(max_len)
            ],
            col_statuses=[
                i < len(t_order) and i < len(s_order) and t_order[i] == s_order[i]
                for i in range(max_len)
            ]
        )
    doc.add_paragraph()


# ==================================================================
# SECTION 3 -- LAYER 2: DATA-LEVEL
# ==================================================================

def _section_layer2(doc, results):
    BLUE = (0x1A, 0x5C, 0x8A)
    l2   = (results.get("layers") or {}).get("layer2")

    _heading(doc, "3. Layer 2 - Data-Level Validation", 1, BLUE)

    if not l2:
        doc.add_paragraph("Layer 2 was skipped due to a Layer 1 structural failure.")
        doc.add_paragraph()
        return

    chks = l2.get("checks", {})
    doc.add_paragraph("Validates row-level correctness using SHA-256 hashing.")
    doc.add_paragraph()

    # 2.1 Row count
    _heading(doc, "3.1 Row Count", 2, BLUE)
    c21 = chks.get("2.1_row_count", {})
    _make_table(doc,
        headers=["Talend Rows", "SQL Rows", "Difference", "Status"],
        rows_data=[(str(c21.get("talend_rows","-")), str(c21.get("sql_rows","-")),
                    str(c21.get("diff", 0)), "PASS" if c21.get("passed") else "FAIL")],
        col_statuses=[c21.get("passed", False)]
    )
    doc.add_paragraph()

    # 2.2 Row hash
    _heading(doc, "3.2 Full Row Hash", 2, BLUE)
    c22 = chks.get("2.2_row_hash", {})
    _make_table(doc,
        headers=["Metric", "Talend", "SQL", "Value"],
        rows_data=[
            ("Total hashes",      str(c22.get("matched_hashes",0) + c22.get("only_in_talend",0)),
                                  str(c22.get("matched_hashes",0) + c22.get("only_in_sql",0)), "-"),
            ("Matched hashes",    "-", "-", str(c22.get("matched_hashes", 0))),
            ("Only in Talend",    str(c22.get("only_in_talend", 0)), "-", "-"),
            ("Only in SQL",       "-", str(c22.get("only_in_sql", 0)), "-"),
            ("Match Rate",        "-", "-", f"{c22.get('match_rate', 0):.2%}"),
        ],
        col_statuses=[c22.get("passed", False)] * 5
    )
    doc.add_paragraph()

    # 2.3 Column hash
    _heading(doc, "3.3 Column-Level Hash", 2, BLUE)
    c23      = chks.get("2.3_column_hash", {})
    diff_cols = c23.get("differing_columns", [])
    clean     = c23.get("clean_columns", [])
    _make_table(doc,
        headers=["Category", "Columns", "Count"],
        rows_data=[
            ("Columns with differences", ", ".join(diff_cols) or "None", str(len(diff_cols))),
            ("Clean columns",            ", ".join(clean) or "None",     str(len(clean))),
        ],
        col_statuses=[len(diff_cols) == 0, True]
    )
    doc.add_paragraph()

    # 2.4 Targeted diff sample
    _heading(doc, "3.4 Targeted Row Diff", 2, BLUE)
    c24    = chks.get("2.4_targeted_diff", {})
    sample = c24.get("sample_diffs", [])
    total  = c24.get("total_diffs", 0)
    if not sample:
        doc.add_paragraph("No cell-level differences found.")
    else:
        doc.add_paragraph(f"Total cell differences: {total}   |   Showing first {len(sample)}")
        doc.add_paragraph()
        _make_table(doc,
            headers=["Row", "Column", "Talend Value", "SQL Value"],
            rows_data=[(str(d["row"]), d["column"], d["talend_value"], d["sql_value"])
                       for d in sample],
            col_statuses=[False] * len(sample)
        )
    doc.add_paragraph()


# ==================================================================
# SECTION 4 -- LAYER 3: BUSINESS RULES
# ==================================================================

def _section_layer3(doc, results):
    BLUE = (0x1A, 0x5C, 0x8A)
    l3   = (results.get("layers") or {}).get("layer3")

    _heading(doc, "4. Layer 3 - Business Rule Validation", 1, BLUE)

    if not l3:
        doc.add_paragraph("Layer 3 was skipped due to a Layer 1 structural failure.")
        doc.add_paragraph()
        return

    chks = l3.get("checks", {})
    doc.add_paragraph(
        "Validates semantic correctness using rules inferred automatically from the data. "
        "No column names or business rules are hardcoded."
    )
    doc.add_paragraph()

    # 3.2 Exact match (inferred categorical columns)
    _heading(doc, "4.1 Exact Match on Inferred Categorical Columns", 2, BLUE)
    exc = chks.get("3.2_exact_match", {})
    if exc:
        _make_table(doc,
            headers=["Column", "Mismatches", "Talend Unique", "SQL Unique", "Status"],
            rows_data=[
                (col,
                 str(res.get("mismatches", 0)),
                 "-",  # cardinality comes from L4
                 "-",
                 "PASS" if res.get("passed") else "FAIL")
                for col, res in exc.items()
            ],
            col_statuses=[res.get("passed", False) for res in exc.values()]
        )
    doc.add_paragraph()

    # 3.3 Relative tolerance
    _heading(doc, "4.2 Numeric Tolerance (Relative) -- Inferred Numeric Columns", 2, BLUE)
    tol = chks.get("3.3_relative_tolerance", {})
    if tol:
        _make_table(doc,
            headers=["Column", "Tolerance", "Mismatches", "Status"],
            rows_data=[
                (col,
                 res.get("tolerance", "-"),
                 str(res.get("mismatches", 0)),
                 "PASS" if res.get("passed") else "FAIL")
                for col, res in tol.items()
            ],
            col_statuses=[res.get("passed", False) for res in tol.values()]
        )
        # Show samples where values differ
        for col, res in tol.items():
            if not res.get("passed") and res.get("sample"):
                doc.add_paragraph()
                doc.add_paragraph(f"Tolerance violations in '{col}':")
                _make_table(doc,
                    headers=["Row", "Talend Value", "SQL Value", "Rel Diff %", "Tolerance %"],
                    rows_data=[
                        (str(s["row"]), str(s["talend_value"]), str(s["sql_value"]),
                         f"{s['rel_diff_pct']}%", f"{s['tolerance_pct']}%")
                        for s in res["sample"]
                    ],
                    col_statuses=[False] * len(res["sample"])
                )
    doc.add_paragraph()

    # 3.4 Aggregation checks
    _heading(doc, "4.3 Aggregation Validation -- Inferred (Numeric x Categorical) Pairs", 2, BLUE)
    agg = chks.get("3.4_aggregations", [])
    if agg:
        _make_table(doc,
            headers=["Check", "Group Column", "Agg Column", "Issues", "Status"],
            rows_data=[
                (a["label"], a["group_col"], a["agg_col"],
                 str(len(a.get("issues", []))),
                 "PASS" if a.get("passed") else "FAIL")
                for a in agg
            ],
            col_statuses=[a.get("passed", False) for a in agg]
        )
        for a in agg:
            if not a.get("passed") and a.get("issues"):
                doc.add_paragraph()
                doc.add_paragraph(f"Issues in '{a['label']}':")
                for issue in a["issues"]:
                    p = doc.add_paragraph(style="List Bullet")
                    p.add_run(issue)
    else:
        doc.add_paragraph("No aggregation checks configured.")
    doc.add_paragraph()

    # 3.5 Referential integrity
    _heading(doc, "4.4 Referential Integrity -- All Columns", 2, BLUE)
    ref = chks.get("3.5_referential_integrity", {})
    if ref:
        for col, res in ref.items():
            _make_table(doc,
                headers=["Column", "Orphans in SQL", "Missing in SQL", "Status"],
                rows_data=[(
                    col,
                    ", ".join(res.get("orphan_in_sql", [])) or "None",
                    ", ".join(res.get("missing_in_sql", [])) or "None",
                    "PASS" if res.get("passed") else "FAIL"
                )],
                col_statuses=[res.get("passed", False)]
            )
    else:
        doc.add_paragraph("No referential integrity checks run.")
    doc.add_paragraph()


# ==================================================================
# SECTION 5 -- LAYER 4: STATISTICAL
# ==================================================================

def _section_layer4(doc, results):
    BLUE = (0x1A, 0x5C, 0x8A)
    l4   = (results.get("layers") or {}).get("layer4")

    _heading(doc, "5. Layer 4 - Statistical Validation", 1, BLUE)

    if not l4:
        doc.add_paragraph("Layer 4 was skipped due to a Layer 1 structural failure.")
        doc.add_paragraph()
        return

    chks      = l4.get("checks", {})
    all_stats = l4.get("all_col_stats", {})
    doc.add_paragraph("Compares statistical distributions between both datasets.")
    doc.add_paragraph()

    # Distribution issues
    _heading(doc, "5.1 Distributions (Min / Max / Mean / Std)", 2, BLUE)
    dist   = chks.get("4.1_4.2_distributions", {})
    issues = dist.get("issues", [])
    if not issues:
        doc.add_paragraph("All numeric distributions match within configured tolerances.")
    else:
        _make_table(doc,
            headers=["Column", "Stat", "Talend", "SQL", "Rel Diff %", "Tolerance %"],
            rows_data=[
                (i["column"], i["stat"], str(i["talend"]), str(i["sql"]),
                 f"{i['rel_diff']}%", f"{i['tol_pct']}%")
                for i in issues
            ],
            col_statuses=[False] * len(issues)
        )
    doc.add_paragraph()

    # Full stats table per numeric column
    _heading(doc, "5.2 Full Statistics per Column (Talend vs SQL)", 2, BLUE)
    numeric_stats = {col: s for col, s in all_stats.items() if s.get("numeric") and "stats" in s}
    if numeric_stats:
        stat_rows = []
        for col, s in numeric_stats.items():
            st = s["stats"]
            for metric in ["min", "max", "mean", "std"]:
                tv = st.get(metric, {}).get("talend", "-")
                sv = st.get(metric, {}).get("sql",    "-")
                match = str(tv) == str(sv)
                stat_rows.append((col, metric.upper(), str(tv), str(sv), "OK" if match else "!!"))
        _make_table(doc,
            headers=["Column", "Stat", "Talend", "SQL", "Match"],
            rows_data=stat_rows,
            col_statuses=[r[4] == "OK" for r in stat_rows]
        )
    else:
        doc.add_paragraph("No numeric columns available for distribution analysis.")
    doc.add_paragraph()

    # Percentiles
    _heading(doc, "5.3 Percentiles (P50 / P95 / P99)", 2, BLUE)
    pct_data = chks.get("4.5_percentiles", {}).get("data", {})
    if pct_data:
        pct_rows = []
        for col, pcts in pct_data.items():
            for pname, vals in pcts.items():
                tv    = vals.get("talend", "-")
                sv    = vals.get("sql",    "-")
                match = str(tv) == str(sv)
                pct_rows.append((col, pname, str(tv), str(sv), "OK" if match else "!!"))
        _make_table(doc,
            headers=["Column", "Percentile", "Talend", "SQL", "Match"],
            rows_data=pct_rows,
            col_statuses=[r[4] == "OK" for r in pct_rows]
        )
    else:
        doc.add_paragraph("No percentile data available.")
    doc.add_paragraph()

    # Null ratio
    _heading(doc, "5.4 Null Ratio per Column", 2, BLUE)
    null_iss = chks.get("4.3_null_ratio", {}).get("issues", [])
    null_rows = []
    for col, s in all_stats.items():
        nr    = s.get("null_ratio", {})
        tv    = nr.get("talend", 0)
        sv    = nr.get("sql",    0)
        diff  = nr.get("diff",   0)
        match = diff == 0
        null_rows.append((col, f"{tv}%", f"{sv}%", f"{diff}%", "OK" if match else "!!"))
    if null_rows:
        _make_table(doc,
            headers=["Column", "Talend Null %", "SQL Null %", "Diff %", "Status"],
            rows_data=null_rows,
            col_statuses=[r[4] == "OK" for r in null_rows]
        )
    doc.add_paragraph()

    # Cardinality
    _heading(doc, "5.5 Cardinality (Unique Value Count)", 2, BLUE)
    card_rows = []
    for col, s in all_stats.items():
        cr    = s.get("cardinality", {})
        tv    = cr.get("talend", "-")
        sv    = cr.get("sql",    "-")
        diff  = cr.get("diff",    0)
        match = diff == 0
        card_rows.append((col, str(tv), str(sv), str(diff), "OK" if match else "!!"))
    if card_rows:
        _make_table(doc,
            headers=["Column", "Talend Unique", "SQL Unique", "Diff", "Status"],
            rows_data=card_rows,
            col_statuses=[r[4] == "OK" for r in card_rows]
        )
    doc.add_paragraph()


# ==================================================================
# SECTION 6 -- ROW-LEVEL DIFFERENCES
# ==================================================================

def _section_diffs(doc, results):
    BLUE = (0x1A, 0x5C, 0x8A)
    rl   = results.get("row_level", {})

    _heading(doc, "6. Row-Level Differences", 1, BLUE)
    _make_table(doc,
        headers=["Metric", "Value"],
        rows_data=[
            ("Total rows compared", str(rl.get("total_rows", "-"))),
            ("Matched rows",        str(rl.get("matched_rows", "-"))),
            ("Diff rows",           str(rl.get("diff_rows", 0))),
            ("Match rate",          f"{rl.get('match_rate', 0):.2%}"),
        ],
        col_statuses=[True, True, rl.get("diff_rows", 0) == 0, rl.get("match_rate",0) >= results["threshold"]]
    )
    doc.add_paragraph()

    sample = rl.get("sample_diffs", [])
    _heading(doc, "6.1 Sample Differences (first 10)", 2, BLUE)
    if not sample:
        doc.add_paragraph("No differences found -- all rows match.")
    else:
        _make_table(doc,
            headers=["Row", "Column", "Talend Value", "SQL Value"],
            rows_data=[(str(d["row"]), d["column"], d["talend_value"], d["sql_value"])
                       for d in sample],
            col_statuses=[False] * len(sample)
        )
    doc.add_paragraph()


# ==================================================================
# SECTION 7 -- LLM REVIEW
# ==================================================================

def _section_llm(doc, results):
    from docx.shared import Pt, RGBColor
    BLUE  = (0x1A, 0x5C, 0x8A)
    GREEN = RGBColor(0x1E, 0x8B, 0x4C)
    RED   = RGBColor(0xC0, 0x39, 0x2B)
    AMBER = RGBColor(0xE6, 0x7E, 0x22)
    GREY  = RGBColor(0x66, 0x66, 0x66)

    _heading(doc, "7. LLM Review (Claude)", 1, BLUE)
    llm = results.get("llm_review")

    if not llm:
        doc.add_paragraph("LLM review not run. Execute run_llm_reviewer.py to add AI analysis.")
        doc.add_paragraph()
        return

    p = doc.add_paragraph()
    p.add_run(f"Mode: {'Claude API (Live)' if llm['mode'] == 'live' else 'Mock Mode'}  |  Reviewed at: {llm.get('reviewed_at', 'N/A')}").italic = True
    doc.add_paragraph()

    llm_dec   = llm.get("final_decision", "N/A")
    dec_color = GREEN if llm_dec == "PASS" else (AMBER if llm_dec == "PASS_WITH_WARNINGS" else RED)
    badge     = doc.add_paragraph()
    r         = badge.add_run(f"  LLM DECISION: {llm_dec}  ")
    r.bold = True; r.font.size = Pt(13); r.font.color.rgb = dec_color
    doc.add_paragraph()

    risk     = llm.get("risk_level", "N/A")
    risk_col = GREEN if "LOW" in str(risk) else (AMBER if "MEDIUM" in str(risk) else RED)
    rp = doc.add_paragraph()
    rr = rp.add_run(f"Risk Level: {risk}")
    rr.bold = True; rr.font.color.rgb = risk_col
    doc.add_paragraph()

    for title, key in [
        ("Decision Reasoning",   "decision_reasoning"),
        ("Difference Analysis",  "difference_analysis"),
        ("SQL Fix",              "sql_fix"),
    ]:
        val = llm.get(key)
        if val:
            _heading(doc, title, 2, BLUE)
            for line in val.splitlines():
                if line.strip():
                    doc.add_paragraph(line.strip())
            doc.add_paragraph()


# ==================================================================
# SECTION 8 -- APPENDIX: FULL DIFF TABLE
# ==================================================================

def _section_appendix(doc, results):
    BLUE = (0x1A, 0x5C, 0x8A)
    rl   = results.get("row_level", {})
    full = rl.get("full_diffs", [])

    _heading(doc, "8. Appendix - Full Difference Table", 1, BLUE)

    if not full:
        doc.add_paragraph("No differences found -- full diff table is empty.")
    else:
        doc.add_paragraph(f"Total cell-level differences: {len(full)}")
        doc.add_paragraph()
        _make_table(doc,
            headers=["Row", "Column", "Talend Value", "SQL Value"],
            rows_data=[(str(d["row"]), d["column"], d["talend_value"], d["sql_value"])
                       for d in full],
            col_statuses=[False] * len(full)
        )
    doc.add_paragraph()


# ==================================================================
# BUILD REPORT
# ==================================================================

def build_report(results: dict) -> str:
    doc = _make_doc()

    _section_cover(doc, results)
    doc.add_page_break()
    _section_executive_summary(doc, results)
    doc.add_page_break()
    _section_layer1(doc, results)
    doc.add_page_break()
    _section_layer2(doc, results)
    doc.add_page_break()
    _section_layer3(doc, results)
    doc.add_page_break()
    _section_layer4(doc, results)
    doc.add_page_break()
    _section_diffs(doc, results)
    doc.add_page_break()
    _section_llm(doc, results)
    doc.add_page_break()
    _section_appendix(doc, results)

    os.makedirs(OUTPUT_DIR, exist_ok=True)
    ts   = datetime.now().strftime("%Y%m%d_%H%M%S")
    path = os.path.join(OUTPUT_DIR, f"validation_report_{ts}.docx")
    doc.save(path)
    return path


# ==================================================================
# MAIN
# ==================================================================

def main():
    print("\n" + SEPARATOR)
    print("  VALIDATION AGENT - REPORTER")
    print(SEPARATOR + "\n")

    print("STEP 1: Running 4-layer comparison")
    print(STEP_LINE)
    try:
        results = run_comparison()
    except Exception as e:
        logger.error(f"Comparison failed: {e}")
        logger.error("Make sure run_executor.py ran successfully first.")
        sys.exit(1)
    print()

    print("STEP 2: LLM Review")
    print(STEP_LINE)
    try:
        llm_mod = _load_module("run_llm_reviewer", "run_llm_reviewer.py")
        results = llm_mod.run_llm_review(results)
    except Exception as e:
        logger.warning(f"LLM review skipped: {e}")
        logger.warning("Report will be generated without LLM section.")
    print()

    print("STEP 3: Generating Word report")
    print(STEP_LINE)
    try:
        report_path = build_report(results)
        logger.success(f"Report saved: {report_path}")
    except Exception as e:
        logger.error(f"Report generation failed: {e}")
        raise
    print()

    print(SEPARATOR)
    comp_dec = results["decision"]
    llm_dec  = results.get("llm_review", {}).get("final_decision")
    risk     = results.get("llm_review", {}).get("risk_level", "N/A")
    final    = llm_dec or comp_dec

    logger.info(f"Comparison decision : {comp_dec}")
    if llm_dec:
        if   llm_dec == "PASS":               logger.success(f"LLM decision        : {llm_dec}")
        elif llm_dec == "PASS_WITH_WARNINGS":  logger.warning(f"LLM decision        : {llm_dec}")
        else:                                  logger.error(f"LLM decision        : {llm_dec}")
    logger.info(f"Risk level          : {risk}")
    logger.info(f"Match rate          : {results['match_rate']:.2%}")
    logger.info(f"Report              : {report_path}")
    print(SEPARATOR + "\n")

    return final in ["PASS", "PASS_WITH_WARNINGS"]


if __name__ == "__main__":
    success = main()
    sys.exit(0 if success else 1)